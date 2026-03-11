import os
from typing import Any, Dict, List

import pandas as pd
import PyPDF2
import pytesseract
from docx import Document
from PIL import Image

def parse_file(filepath: str) -> Dict[str, Any]:
    """解析不同格式文件，并统一返回结构。"""
    if not filepath or not os.path.exists(filepath):
        raise FileNotFoundError(f'File does not exist: {filepath}')

    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.pdf':
        return parse_pdf(filepath)
    if ext == '.docx':
        return parse_docx(filepath)
    if ext == '.doc':
        raise Exception('Legacy .doc is not supported. Please convert to .docx and retry.')
    if ext in ['.txt', '.log']:
        return parse_txt(filepath)
    if ext == '.md':
        return parse_md(filepath)
    if ext in ['.xlsx', '.xls']:
        return parse_excel(filepath)
    if ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tif', '.tiff']:
        return parse_image(filepath)

    raise Exception(f'Unsupported file format: {ext}')

def parse_pdf(filepath: str) -> Dict[str, Any]:
    """解析PDF文件。"""
    text_chunks: List[str] = []
    images: List[str] = []
    tables: List[Any] = []

    try:
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)

            if getattr(reader, 'is_encrypted', False):
                try:
                    reader.decrypt('')
                except Exception:
                    raise Exception('Encrypted PDF is not supported. Please provide an unencrypted file.')

            num_pages = len(reader.pages)
            print(f"PDF页数: {num_pages}")
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text() or ''
                print(f"第{page_num+1}页文本长度: {len(page_text)}")
                if page_text:
                    text_chunks.append(page_text)
                    print(f"第{page_num+1}页前100个字符: {page_text[:100]}...")

        return {
            'text': '\n'.join(text_chunks).strip(),
            'images': images,
            'tables': tables,
            'page_count': num_pages,
        }
    except Exception as e:
        raise Exception(f'Failed to parse PDF: {str(e)}')

def parse_docx(filepath: str) -> Dict[str, Any]:
    """解析Word(docx)文件。"""
    images: List[str] = []
    tables: List[List[List[str]]] = []

    try:
        doc = Document(filepath)

        text_lines = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() if cell.text else '' for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append(table_data)

        return {
            'text': '\n'.join(text_lines),
            'images': images,
            'tables': tables,
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
        }
    except Exception as e:
        raise Exception(f'Failed to parse Word document: {str(e)}')

def _read_text_with_fallback(filepath: str) -> str:
    encodings = ['utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'big5', 'latin-1']
    for enc in encodings:
        try:
            with open(filepath, 'r', encoding=enc, errors='strict') as f:
                return f.read()
        except Exception:
            continue

    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def parse_txt(filepath: str) -> Dict[str, Any]:
    """解析文本文件。"""
    text = _read_text_with_fallback(filepath)
    return {
        'text': text,
        'line_count': len(text.splitlines())
    }

def parse_md(filepath: str) -> Dict[str, Any]:
    """解析Markdown文件。"""
    text = _read_text_with_fallback(filepath)
    return {
        'text': text,
        'line_count': len(text.splitlines())
    }

def parse_excel(filepath: str) -> Dict[str, Any]:
    """解析Excel文件。"""
    tables = []
    ext = os.path.splitext(filepath)[1].lower()

    try:
        engine = 'xlrd' if ext == '.xls' else 'openpyxl'
        xl = pd.ExcelFile(filepath, engine=engine)
    except Exception as e:
        raise Exception(f'Failed to open Excel file: {str(e)}')

    for sheet_name in xl.sheet_names:
        try:
            df = xl.parse(sheet_name)
        except Exception:
            continue

        if df.empty:
            continue

        df = df.fillna('')
        table_data = df.values.tolist()
        headers = [str(h) for h in df.columns.tolist()]
        table_data.insert(0, headers)
        tables.append({'sheet_name': sheet_name, 'data': table_data})

    return {
        'tables': tables,
        'sheet_count': len(tables)
    }


def parse_image(filepath: str) -> Dict[str, Any]:
    """解析图片文件，提取OCR文本。"""
    text = ocr_image(filepath)
    return {
        'text': text,
        'images': [filepath],
        'tables': [],
    }

def ocr_image(image_path: str) -> str:
    """OCR识别图片中的文字。"""
    try:
        image = Image.open(image_path)
        return pytesseract.image_to_string(image, lang='chi_sim')
    except Exception:
        return ''